"""Router — Module Documents IA.

Préfixe : /ai-documents
Module requis : ai_documents
"""

import asyncio
import io
import json
import logging
import re
from typing import AsyncGenerator

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, status
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from sqlalchemy import or_

from app.database import get_db
from app.deps import get_current_user, require_module, require_super_admin
from app.models.ai_documents import AIDocument, AIDocumentTemplate, template_tenant_assignments
from app.models.user import User
from app.schemas.ai_documents import (
    AIDocumentListItem,
    AIDocumentResponse,
    GenerateRequest,
    TemplateAssignRequest,
    TemplateCreate,
    TemplateResponse,
    TemplateUpdate,
)
from app.models.transcription import TranscriptionJob
from app.services.ai_documents import enqueue_generation
from app.services.event_bus import event_bus

logger = logging.getLogger(__name__)

router = APIRouter(
    prefix="/ai-documents",
    tags=["AI Documents"],
    dependencies=[Depends(require_module("ai_documents"))],
)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _template_to_response(tpl: AIDocumentTemplate) -> dict:
    """Convert a template ORM object to a response dict with assigned_tenant_ids."""
    data = {c.key: getattr(tpl, c.key) for c in AIDocumentTemplate.__table__.columns}
    data["assigned_tenant_ids"] = [t.id for t in tpl.assigned_tenants] if tpl.is_global else []
    # Parse workflow_steps JSON → list
    if data.get("workflow_steps") and isinstance(data["workflow_steps"], str):
        try:
            data["workflow_steps"] = json.loads(data["workflow_steps"])
        except (json.JSONDecodeError, TypeError):
            data["workflow_steps"] = None
    return data


def _get_template_or_404(template_id: str, tenant_id: str, db: Session) -> AIDocumentTemplate:
    """Get a template owned by tenant OR a global template assigned to tenant."""
    tpl = db.query(AIDocumentTemplate).filter_by(id=template_id, tenant_id=tenant_id).first()
    if not tpl:
        # Also check global templates assigned to this tenant
        tpl = (
            db.query(AIDocumentTemplate)
            .join(template_tenant_assignments)
            .filter(
                AIDocumentTemplate.id == template_id,
                AIDocumentTemplate.is_global == True,
                template_tenant_assignments.c.tenant_id == tenant_id,
            )
            .first()
        )
    if not tpl:
        raise HTTPException(status_code=404, detail="Template introuvable")
    return tpl


def _get_document_or_404(doc_id: str, tenant_id: str, db: Session) -> AIDocument:
    doc = db.query(AIDocument).filter_by(id=doc_id, tenant_id=tenant_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document introuvable")
    return doc


# ── Templates ─────────────────────────────────────────────────────────────────

@router.get("/templates", response_model=list[TemplateResponse])
def list_templates(
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Liste les templates du tenant + les templates globaux assignés."""
    # Tenant-owned templates
    owned = (
        db.query(AIDocumentTemplate)
        .filter_by(tenant_id=user.tenant_id)
        .order_by(AIDocumentTemplate.created_at)
        .all()
    )
    # Global templates assigned to this tenant
    assigned = (
        db.query(AIDocumentTemplate)
        .join(template_tenant_assignments)
        .filter(
            AIDocumentTemplate.is_global == True,
            template_tenant_assignments.c.tenant_id == user.tenant_id,
        )
        .order_by(AIDocumentTemplate.created_at)
        .all()
    )
    # Merge and deduplicate
    seen = set()
    result = []
    for tpl in owned + assigned:
        if tpl.id not in seen:
            seen.add(tpl.id)
            result.append(_template_to_response(tpl))
    return result


@router.post("/templates", response_model=TemplateResponse, status_code=201)
def create_template(
    body: TemplateCreate,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    data = body.model_dump()
    # Serialize workflow_steps to JSON string for storage
    if data.get("workflow_steps") is not None:
        data["workflow_steps"] = json.dumps([s.model_dump() for s in body.workflow_steps], ensure_ascii=False) if body.workflow_steps else None
    tpl = AIDocumentTemplate(tenant_id=user.tenant_id, **data)
    db.add(tpl)
    db.commit()
    db.refresh(tpl)
    return _template_to_response(tpl)


_EXPORT_FIELDS = [
    "name", "description", "document_type", "system_prompt",
    "user_prompt_template", "map_system_prompt", "ollama_model",
    "temperature", "is_active",
]


@router.get("/templates/export")
def export_templates(
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Exporte tous les templates du tenant en JSON."""
    templates = (
        db.query(AIDocumentTemplate)
        .filter_by(tenant_id=user.tenant_id)
        .order_by(AIDocumentTemplate.created_at)
        .all()
    )
    payload = [
        {f: getattr(t, f) for f in _EXPORT_FIELDS}
        for t in templates
    ]
    content = json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
    return StreamingResponse(
        io.BytesIO(content),
        media_type="application/json",
        headers={"Content-Disposition": "attachment; filename=templates_export.json"},
    )


@router.post("/templates/import")
def import_templates(
    file: UploadFile = File(...),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Importe des templates depuis un fichier JSON."""
    try:
        data = json.loads(file.file.read())
    except (json.JSONDecodeError, UnicodeDecodeError):
        raise HTTPException(status_code=400, detail="Fichier JSON invalide")

    if not isinstance(data, list):
        raise HTTPException(status_code=400, detail="Le fichier doit contenir un tableau de templates")

    created = 0
    errors = []
    for idx, item in enumerate(data):
        name = item.get("name", "").strip()
        system_prompt = item.get("system_prompt", "").strip()
        user_prompt_template = item.get("user_prompt_template", "").strip()
        if not name or not system_prompt or not user_prompt_template:
            errors.append(f"Template #{idx + 1} : champs obligatoires manquants (name, system_prompt, user_prompt_template)")
            continue
        tpl = AIDocumentTemplate(
            tenant_id=user.tenant_id,
            name=name,
            description=item.get("description"),
            document_type=item.get("document_type", "custom"),
            system_prompt=system_prompt,
            user_prompt_template=user_prompt_template,
            map_system_prompt=item.get("map_system_prompt"),
            ollama_model=item.get("ollama_model"),
            temperature=item.get("temperature", 0.3),
            is_active=item.get("is_active", True),
        )
        db.add(tpl)
        created += 1
    db.commit()
    return {"created": created, "errors": errors}


@router.get("/templates/{template_id}", response_model=TemplateResponse)
def get_template(
    template_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    tpl = _get_template_or_404(template_id, user.tenant_id, db)
    return _template_to_response(tpl)


@router.patch("/templates/{template_id}", response_model=TemplateResponse)
def update_template(
    template_id: str,
    body: TemplateUpdate,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    tpl = _get_template_or_404(template_id, user.tenant_id, db)
    if tpl.is_global:
        raise HTTPException(status_code=403, detail="Les templates globaux ne peuvent être modifiés que par un super admin")
    for field, value in body.model_dump(exclude_unset=True).items():
        if field == "workflow_steps" and value is not None:
            value = json.dumps(value, ensure_ascii=False)
        setattr(tpl, field, value)
    db.commit()
    db.refresh(tpl)
    return _template_to_response(tpl)


@router.delete("/templates/{template_id}", status_code=204)
def delete_template(
    template_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    from app.models.procedures import ProcedureTemplate, Procedure
    tpl = _get_template_or_404(template_id, user.tenant_id, db)
    # Nettoyer les FK avant suppression
    db.query(AIDocument).filter(AIDocument.template_id == template_id).update(
        {AIDocument.template_id: None}, synchronize_session=False)
    db.query(ProcedureTemplate).filter(ProcedureTemplate.document_template_id == template_id).update(
        {ProcedureTemplate.document_template_id: None}, synchronize_session=False)
    db.query(Procedure).filter(Procedure.document_template_id == template_id).update(
        {Procedure.document_template_id: None}, synchronize_session=False)
    db.delete(tpl)
    db.commit()


# ── Workflow generation for procedure templates ─────────────────────────────

_WORKFLOW_SYSTEM_PROMPT = """Tu es un expert en modélisation de workflows et procédures métier.
On te donne un template de procédure. Tu dois produire les ÉTAPES SÉQUENTIELLES du workflow.

Les étapes utilisent les modules existants de l'application :
- "form" : formulaire à remplir par l'utilisateur (config.fields = [{id, label, type, required}])
  type de champ : "text", "textarea", "date", "file"
- "select_contacts" : sélection de contacts/participants depuis le module Contacts (config = {})
- "send_email" : envoi d'email aux contacts sélectionnés (config.subject_template, config.body_template)
- "collect_responses" : envoi de formulaires aux participants et attente des réponses
  (config.roles = [{role_name, invitation_delay_days, form_questions: [{id, label, type, required}]}])
- "generate_document" : génération d'un document IA via le module Documents (config = {})
- "upload_document" : upload d'un fichier (config.accepted_types = ["pdf","docx"])
- "manual" : étape manuelle / validation humaine (config.instructions = "...")

Retourne UNIQUEMENT un JSON valide (un tableau d'étapes), sans commentaire, sans markdown.
[{"step_type": "...", "label": "...", "description": "...", "config": {...}, "is_required": true}]"""


@router.post("/templates/{template_id}/generate-workflow")
def generate_template_workflow(
    template_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Génère le workflow via LLM pour un template de catégorie 'procedure'."""
    import uuid
    from app.services.llm_client import llm_generate, resolve_model
    from app.services.ai_config import get_model_for_usage

    tpl = _get_template_or_404(template_id, user.tenant_id, db)
    if tpl.category != "procedure":
        raise HTTPException(status_code=400, detail="Ce template n'est pas de catégorie 'procedure'")

    model = resolve_model(get_model_for_usage("workflow_generation"))
    user_prompt = (
        f"Nom de la procédure : {tpl.name}\n"
        f"Description : {tpl.description or 'Aucune'}\n\n"
    )
    if tpl.user_prompt_template:
        user_prompt += f"Contexte supplémentaire :\n{tpl.user_prompt_template}\n\n"
    user_prompt += "Génère les étapes séquentielles de ce workflow."

    try:
        raw = llm_generate(
            model=model,
            system_prompt=_WORKFLOW_SYSTEM_PROMPT,
            user_prompt=user_prompt,
            temperature=0.3,
            extra_params={"keep_alive": 0},
        )
    except RuntimeError as exc:
        raise HTTPException(status_code=502, detail=f"Erreur LLM : {exc}")

    # Parse JSON (handle markdown code blocks)
    cleaned = raw.strip()
    if cleaned.startswith("```"):
        lines = cleaned.split("\n")
        lines = [l for l in lines if not l.strip().startswith("```")]
        cleaned = "\n".join(lines).strip()

    try:
        data = json.loads(cleaned)
    except json.JSONDecodeError:
        raise HTTPException(status_code=422, detail="Le LLM n'a pas retourné un JSON valide. Réessayez.")

    # Normaliser : tableau d'étapes ou objet avec "steps"
    if isinstance(data, dict):
        steps = data.get("steps", [])
    elif isinstance(data, list):
        if data and isinstance(data[0], dict) and "steps" in data[0]:
            steps = data[0]["steps"]
        else:
            steps = data
    else:
        steps = []

    # Assigner des IDs uniques aux champs de formulaire
    for step in steps:
        config = step.get("config", {}) or {}
        for i, f in enumerate(config.get("fields", [])):
            if not f.get("id"):
                f["id"] = f"f{i+1}_{uuid.uuid4().hex[:6]}"
        for role in config.get("roles", []):
            for i, q in enumerate(role.get("form_questions", [])):
                if not q.get("id"):
                    q["id"] = f"q{i+1}_{uuid.uuid4().hex[:6]}"

    # Sauvegarder dans le template
    tpl.workflow_steps = json.dumps(steps, ensure_ascii=False)
    db.commit()
    db.refresh(tpl)

    return _template_to_response(tpl)


@router.post("/global-templates/{template_id}/generate-workflow")
def generate_global_template_workflow(
    template_id: str,
    _: User = Depends(require_super_admin),
    db: Session = Depends(get_db),
):
    """Génère le workflow via LLM pour un template global de catégorie 'procedure'."""
    import uuid
    from app.services.llm_client import llm_generate, resolve_model
    from app.services.ai_config import get_model_for_usage

    tpl = db.query(AIDocumentTemplate).filter_by(id=template_id, is_global=True).first()
    if not tpl:
        raise HTTPException(status_code=404, detail="Template global introuvable")
    if tpl.category != "procedure":
        raise HTTPException(status_code=400, detail="Ce template n'est pas de catégorie 'procedure'")

    model = resolve_model(get_model_for_usage("workflow_generation"))
    user_prompt = (
        f"Nom de la procédure : {tpl.name}\n"
        f"Description : {tpl.description or 'Aucune'}\n\n"
    )
    if tpl.user_prompt_template:
        user_prompt += f"Contexte supplémentaire :\n{tpl.user_prompt_template}\n\n"
    user_prompt += "Génère les étapes séquentielles de ce workflow."

    try:
        raw = llm_generate(
            model=model,
            system_prompt=_WORKFLOW_SYSTEM_PROMPT,
            user_prompt=user_prompt,
            temperature=0.3,
            extra_params={"keep_alive": 0},
        )
    except RuntimeError as exc:
        raise HTTPException(status_code=502, detail=f"Erreur LLM : {exc}")

    cleaned = raw.strip()
    if cleaned.startswith("```"):
        lines = cleaned.split("\n")
        lines = [l for l in lines if not l.strip().startswith("```")]
        cleaned = "\n".join(lines).strip()

    try:
        data = json.loads(cleaned)
    except json.JSONDecodeError:
        raise HTTPException(status_code=422, detail="Le LLM n'a pas retourné un JSON valide. Réessayez.")

    if isinstance(data, dict):
        steps = data.get("steps", [])
    elif isinstance(data, list):
        if data and isinstance(data[0], dict) and "steps" in data[0]:
            steps = data[0]["steps"]
        else:
            steps = data
    else:
        steps = []

    for step in steps:
        config = step.get("config", {}) or {}
        for i, f in enumerate(config.get("fields", [])):
            if not f.get("id"):
                f["id"] = f"f{i+1}_{uuid.uuid4().hex[:6]}"
        for role in config.get("roles", []):
            for i, q in enumerate(role.get("form_questions", [])):
                if not q.get("id"):
                    q["id"] = f"q{i+1}_{uuid.uuid4().hex[:6]}"

    tpl.workflow_steps = json.dumps(steps, ensure_ascii=False)
    db.commit()
    db.refresh(tpl)

    return _template_to_response(tpl)


# ── Global templates (super_admin) ───────────────────────────────────────────

@router.get("/global-templates", response_model=list[TemplateResponse])
def list_global_templates(
    _: User = Depends(require_super_admin),
    db: Session = Depends(get_db),
):
    """Liste tous les templates globaux (super_admin)."""
    templates = (
        db.query(AIDocumentTemplate)
        .filter_by(is_global=True)
        .order_by(AIDocumentTemplate.category, AIDocumentTemplate.created_at)
        .all()
    )
    return [_template_to_response(tpl) for tpl in templates]


@router.post("/global-templates", response_model=TemplateResponse, status_code=201)
def create_global_template(
    body: TemplateCreate,
    _: User = Depends(require_super_admin),
    db: Session = Depends(get_db),
):
    """Crée un template global (super_admin)."""
    data = body.model_dump(exclude={"is_global"})
    if data.get("workflow_steps") is not None:
        data["workflow_steps"] = json.dumps([s.model_dump() for s in body.workflow_steps], ensure_ascii=False) if body.workflow_steps else None
    tpl = AIDocumentTemplate(
        tenant_id=None,
        is_global=True,
        **data,
    )
    db.add(tpl)
    db.commit()
    db.refresh(tpl)
    return _template_to_response(tpl)


@router.get("/global-templates/{template_id}", response_model=TemplateResponse)
def get_global_template(
    template_id: str,
    _: User = Depends(require_super_admin),
    db: Session = Depends(get_db),
):
    tpl = db.query(AIDocumentTemplate).filter_by(id=template_id, is_global=True).first()
    if not tpl:
        raise HTTPException(status_code=404, detail="Template global introuvable")
    return _template_to_response(tpl)


@router.patch("/global-templates/{template_id}", response_model=TemplateResponse)
def update_global_template(
    template_id: str,
    body: TemplateUpdate,
    _: User = Depends(require_super_admin),
    db: Session = Depends(get_db),
):
    tpl = db.query(AIDocumentTemplate).filter_by(id=template_id, is_global=True).first()
    if not tpl:
        raise HTTPException(status_code=404, detail="Template global introuvable")
    for field, value in body.model_dump(exclude_unset=True).items():
        if field == "workflow_steps" and value is not None:
            value = json.dumps(value, ensure_ascii=False)
        setattr(tpl, field, value)
    db.commit()
    db.refresh(tpl)
    return _template_to_response(tpl)


@router.delete("/global-templates/{template_id}", status_code=204)
def delete_global_template(
    template_id: str,
    _: User = Depends(require_super_admin),
    db: Session = Depends(get_db),
):
    from app.models.procedures import ProcedureTemplate, Procedure
    tpl = db.query(AIDocumentTemplate).filter_by(id=template_id, is_global=True).first()
    if not tpl:
        raise HTTPException(status_code=404, detail="Template global introuvable")
    db.query(AIDocument).filter(AIDocument.template_id == template_id).update(
        {AIDocument.template_id: None}, synchronize_session=False)
    db.query(ProcedureTemplate).filter(ProcedureTemplate.document_template_id == template_id).update(
        {ProcedureTemplate.document_template_id: None}, synchronize_session=False)
    db.query(Procedure).filter(Procedure.document_template_id == template_id).update(
        {Procedure.document_template_id: None}, synchronize_session=False)
    db.execute(template_tenant_assignments.delete().where(
        template_tenant_assignments.c.template_id == template_id))
    db.delete(tpl)
    db.commit()


@router.put("/global-templates/{template_id}/assign", response_model=TemplateResponse)
def assign_global_template(
    template_id: str,
    body: TemplateAssignRequest,
    _: User = Depends(require_super_admin),
    db: Session = Depends(get_db),
):
    """Assign a global template to specific tenants (replaces previous assignments)."""
    tpl = db.query(AIDocumentTemplate).filter_by(id=template_id, is_global=True).first()
    if not tpl:
        raise HTTPException(status_code=404, detail="Template global introuvable")

    # Clear existing assignments
    db.execute(template_tenant_assignments.delete().where(
        template_tenant_assignments.c.template_id == template_id))

    # Insert new assignments
    from app.models.tenant import Tenant
    for tid in body.tenant_ids:
        tenant = db.query(Tenant).filter_by(id=tid).first()
        if tenant:
            db.execute(template_tenant_assignments.insert().values(
                template_id=template_id, tenant_id=tid))

    db.commit()
    db.refresh(tpl)
    return _template_to_response(tpl)


# ── Modèles Ollama disponibles ────────────────────────────────────────────────

@router.get("/ollama-models")
def list_ollama_models(db: Session = Depends(get_db)):
    """Retourne les modèles disponibles (Ollama local + cloud si activé)."""
    import requests as req
    from app.config import settings
    from app.models.ai_settings import CloudProvider, OVH_MODELS

    models = []
    try:
        resp = req.get(f"{settings.ollama_url}/api/tags", timeout=5)
        resp.raise_for_status()
        data = resp.json()
        models = [m["name"] for m in data.get("models", [])]
    except Exception as exc:
        logger.warning(f"[AI] Impossible de lister les modèles Ollama : {exc}")

    # Append cloud models if enabled
    cloud_provider = db.query(CloudProvider).filter_by(provider_name="ovh", enabled=True).first()
    if cloud_provider:
        for m in OVH_MODELS:
            models.append(f"cloud/{m['id']}")

    return {"models": models, "default": settings.ollama_default_model}


@router.delete("/ollama-models/{model_name:path}", status_code=204)
def delete_ollama_model(
    model_name: str,
    _: User = Depends(require_super_admin),
):
    """Supprime un modèle Ollama."""
    import requests as req
    from app.config import settings
    try:
        resp = req.delete(f"{settings.ollama_url}/api/delete", json={"name": model_name}, timeout=30)
        resp.raise_for_status()
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"Erreur Ollama : {exc}")


@router.get("/ollama-models/pull")
async def pull_ollama_model(
    model: str,
    _: User = Depends(require_super_admin),
):
    """Lance le téléchargement d'un modèle Ollama et stream la progression en SSE."""
    import requests as req
    from app.config import settings

    async def _stream() -> AsyncGenerator[str, None]:
        try:
            with req.post(
                f"{settings.ollama_url}/api/pull",
                json={"name": model, "stream": True},
                stream=True,
                timeout=3600,
            ) as resp:
                resp.raise_for_status()
                for line in resp.iter_lines():
                    if not line:
                        continue
                    yield f"data: {line.decode()}\n\n"
                    try:
                        chunk = json.loads(line)
                        if chunk.get("status") == "success":
                            break
                    except json.JSONDecodeError:
                        pass
        except Exception as exc:
            yield f"data: {json.dumps({'error': str(exc)})}\n\n"

    return StreamingResponse(_stream(), media_type="text/event-stream")


# ── Sources disponibles ───────────────────────────────────────────────────────

@router.get("/sources/sessions")
def list_available_sessions(
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Liste toutes les sessions de transcription (simple + diarisation) completees."""
    jobs = (
        db.query(TranscriptionJob)
        .filter(
            TranscriptionJob.tenant_id == user.tenant_id,
            TranscriptionJob.status == "completed",
        )
        .order_by(TranscriptionJob.created_at.desc())
        .all()
    )
    return [
        {
            "id": j.id,
            "title": j.title,
            "original_filename": j.original_filename,
            "created_at": j.created_at.isoformat() if j.created_at else None,
            "mode": j.mode,
        }
        for j in jobs
    ]


# ── Génération ────────────────────────────────────────────────────────────────

@router.post("/generate", response_model=AIDocumentResponse, status_code=202)
def generate_document(
    body: GenerateRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Lance une génération asynchrone. Retourne immédiatement le document en statut 'pending'."""
    tpl = _get_template_or_404(body.template_id, user.tenant_id, db)

    # ── Consent guard-fou ─────────────────────────────────────────────────
    if body.source_session_id:
        from app.models.transcription import TranscriptionJob
        session = db.query(TranscriptionJob).filter_by(id=body.source_session_id).first()
        if session:
            rv = session.recording_validity
            if rv == "invalidated":
                raise HTTPException(403, "Session invalidée (retrait de consentement). Génération bloquée.")
            if rv == "blocked":
                raise HTTPException(403, "Session bloquée (aucun consentement détecté). Génération impossible.")
            if rv == "pending":
                raise HTTPException(403, "Consentements en attente de validation. Finalisez les consentements avant de générer.")

            # Check individual attendees
            if session.attendees:
                try:
                    attendees = json.loads(session.attendees)
                    statuses = {a.get("status") for a in attendees}
                    if "refused" in statuses or "withdrawn" in statuses:
                        raise HTTPException(403, "Un ou plusieurs participants ont refusé ou retiré leur consentement.")
                    if "pending" in statuses or "pending_oral" in statuses:
                        raise HTTPException(403, "Tous les participants doivent avoir consenti avant la génération.")
                except (json.JSONDecodeError, TypeError):
                    pass

    # Snapshot du template + consent audit trail
    snapshot = {
        "name": tpl.name,
        "document_type": tpl.document_type,
        "system_prompt": tpl.system_prompt,
        "user_prompt_template": tpl.user_prompt_template,
        "map_system_prompt": tpl.map_system_prompt,
        "ollama_model": tpl.ollama_model,
        "temperature": tpl.temperature,
    }

    # Add consent audit trail
    if body.source_session_id:
        from app.models.transcription import TranscriptionJob
        session = db.query(TranscriptionJob).filter_by(id=body.source_session_id).first()
        if session:
            consent_audit = {
                "recording_validity": session.recording_validity,
                "generated_by": user.id,
                "generated_by_name": user.display_name,
            }
            if session.attendees:
                try:
                    consent_audit["attendees"] = json.loads(session.attendees)
                except (json.JSONDecodeError, TypeError):
                    pass
            snapshot["consent_audit"] = consent_audit

    # Auto-resolve dossier from planned meeting if not explicitly provided
    effective_dossier_id = body.source_dossier_id
    if not effective_dossier_id and body.source_session_id:
        from app.models.planned_meeting import PlannedMeeting
        from app.models.preparatory import PreparatoryDossier
        planned = db.query(PlannedMeeting).filter_by(job_id=body.source_session_id).first()
        if planned:
            dossier = db.query(PreparatoryDossier).filter_by(planned_meeting_id=planned.id).first()
            if dossier:
                effective_dossier_id = dossier.id

    doc = AIDocument(
        tenant_id=user.tenant_id,
        user_id=user.id,
        template_id=tpl.id,
        template_snapshot=json.dumps(snapshot, ensure_ascii=False),
        title=body.title,
        status="pending",
        source_dossier_id=effective_dossier_id,
        source_session_id=body.source_session_id,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    enqueue_generation(doc.id)
    logger.info(f"[AI] Document {doc.id} mis en file d'attente")
    return doc


# ── Liste / Détail / Suppression ──────────────────────────────────────────────

@router.get("/documents", response_model=list[AIDocumentListItem])
def list_documents(
    include_invalidated: bool = False,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    q = db.query(AIDocument).filter_by(tenant_id=user.tenant_id)
    if not include_invalidated:
        q = q.filter(AIDocument.invalidated_at.is_(None))
    return q.order_by(AIDocument.created_at.desc()).all()


@router.get("/documents/{doc_id}", response_model=AIDocumentResponse)
def get_document(
    doc_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    return _get_document_or_404(doc_id, user.tenant_id, db)


@router.patch("/documents/{doc_id}", response_model=AIDocumentResponse)
def update_document(
    doc_id: str,
    body: dict,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Met à jour le contenu du document (result_text après édition)."""
    doc = _get_document_or_404(doc_id, user.tenant_id, db)
    if "result_text" in body:
        doc.result_text = body["result_text"]
    db.commit()
    db.refresh(doc)
    return doc


@router.delete("/documents/{doc_id}", status_code=204)
def delete_document(
    doc_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    doc = _get_document_or_404(doc_id, user.tenant_id, db)
    db.delete(doc)
    db.commit()


# ── Export ─────────────────────────────────────────────────────────────────────

def _markdown_to_pdf(title: str, text: str) -> bytes:
    """Convertit un texte Markdown en PDF via fpdf2."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_margins(20, 20, 20)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Titre du document
    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(0, 10, title)
    pdf.ln(4)
    pdf.set_draw_color(180, 180, 180)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(6)

    for line in text.split("\n"):
        stripped = line.strip()

        if stripped.startswith("### "):
            pdf.set_font("Helvetica", "B", 11)
            pdf.multi_cell(0, 7, stripped[4:])
            pdf.ln(1)
        elif stripped.startswith("## "):
            pdf.set_font("Helvetica", "B", 13)
            pdf.multi_cell(0, 8, stripped[3:])
            pdf.ln(2)
        elif stripped.startswith("# "):
            pdf.set_font("Helvetica", "B", 15)
            pdf.multi_cell(0, 9, stripped[2:])
            pdf.ln(3)
        elif stripped in ("---", "***", "___"):
            pdf.ln(2)
            pdf.set_draw_color(200, 200, 200)
            pdf.line(20, pdf.get_y(), 190, pdf.get_y())
            pdf.ln(3)
        elif stripped == "":
            pdf.ln(4)
        else:
            # Retirer le formatage Markdown inline (**bold**, *italic*, `code`)
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            clean = re.sub(r"\*(.+?)\*", r"\1", clean)
            clean = re.sub(r"`(.+?)`", r"\1", clean)

            if stripped.startswith(("- ", "* ", "• ")):
                pdf.set_font("Helvetica", size=10)
                pdf.set_x(25)
                pdf.multi_cell(0, 6, "\u2022 " + clean[2:])
            else:
                pdf.set_font("Helvetica", size=10)
                pdf.multi_cell(0, 6, clean)

    return bytes(pdf.output())


def _text_to_docx(title: str, text: str) -> bytes:
    """Convertit un texte Markdown/HTML en DOCX via python-docx."""
    import io
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(title, level=0)

    # Handle both HTML (from editor save) and Markdown (from AI generation)
    is_html = "<p>" in text or "<h" in text

    if is_html:
        # Strip HTML tags for DOCX (basic approach)
        import html as html_mod
        clean = re.sub(r"<br\s*/?>", "\n", text)
        clean = re.sub(r"</p>", "\n", clean)
        clean = re.sub(r"</?(h[1-6])[^>]*>", "\n", clean)
        clean = re.sub(r"</?(ul|ol|li|div|blockquote)[^>]*>", "\n", clean)
        clean = re.sub(r"<[^>]+>", "", clean)
        clean = html_mod.unescape(clean)
        lines = clean.split("\n")
    else:
        lines = text.split("\n")

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        if not is_html and stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif not is_html and stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif not is_html and stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif not is_html and stripped.startswith(("- ", "* ", "\u2022 ")):
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped[2:])
            doc.add_paragraph(clean, style="List Bullet")
        elif not is_html and stripped in ("---", "***", "___"):
            doc.add_paragraph("_" * 50)
        else:
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            clean = re.sub(r"\*(.+?)\*", r"\1", clean)
            clean = re.sub(r"`(.+?)`", r"\1", clean)
            p = doc.add_paragraph(clean)
            p.style.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.get("/documents/{doc_id}/export")
def export_document(
    doc_id: str,
    format: str = "md",
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Télécharge le document généré (md, txt ou pdf)."""
    doc = _get_document_or_404(doc_id, user.tenant_id, db)
    if doc.invalidated_at:
        raise HTTPException(status_code=403, detail=f"Document invalidé : {doc.invalidated_reason or 'retrait de consentement'}")
    if doc.status != "completed" or not doc.result_text:
        raise HTTPException(status_code=409, detail="Document pas encore disponible")

    safe_title = doc.title.replace("/", "-").replace("\\", "-")

    if format == "pdf":
        try:
            pdf_bytes = _markdown_to_pdf(doc.title, doc.result_text)
        except Exception as exc:
            logger.error(f"[PDF] Erreur génération PDF pour {doc_id}: {exc}")
            raise HTTPException(status_code=500, detail="Erreur lors de la génération du PDF")
        return StreamingResponse(
            iter([pdf_bytes]),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.pdf"'},
        )

    if format == "docx":
        try:
            docx_bytes = _text_to_docx(doc.title, doc.result_text)
        except Exception as exc:
            logger.error(f"[DOCX] Erreur génération DOCX pour {doc_id}: {exc}")
            raise HTTPException(status_code=500, detail="Erreur lors de la génération du DOCX")
        return StreamingResponse(
            iter([docx_bytes]),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.docx"'},
        )

    ext = "md" if format == "md" else "txt"
    return StreamingResponse(
        iter([doc.result_text.encode("utf-8")]),
        media_type="text/plain; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{safe_title}.{ext}"'},
    )


# ── SSE ───────────────────────────────────────────────────────────────────────

@router.get("/documents/{doc_id}/events")
async def document_events(
    doc_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Stream SSE des événements de génération."""
    _get_document_or_404(doc_id, user.tenant_id, db)

    async def _stream() -> AsyncGenerator[str, None]:
        q = event_bus.subscribe(doc_id)
        try:
            while True:
                try:
                    event = await asyncio.wait_for(q.get(), timeout=30.0)
                    yield f"data: {json.dumps(event)}\n\n"
                    if event.get("status") in ("completed", "error"):
                        break
                except asyncio.TimeoutError:
                    yield ": keepalive\n\n"
        finally:
            event_bus.unsubscribe(doc_id, q)

    return StreamingResponse(_stream(), media_type="text/event-stream")
